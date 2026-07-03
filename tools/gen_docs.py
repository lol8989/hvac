# -*- coding: utf-8 -*-
"""도면검도시스템 정책서·WBS(.docx) + 지라 import(.csv) 생성기.
근거: 프로젝트 코드(도메인/유즈케이스), CLAUDE.md, doc/ 요구사항(Function Inventory 등).
실행: python tools/gen_docs.py  → POC 루트에 문서 3종 생성."""
import csv
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FONT = "Malgun Gothic"


def set_ko_font(doc):
    for name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        style.font.name = FONT
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), FONT)
        rfonts.set(qn("w:ascii"), FONT)
        rfonts.set(qn("w:hAnsi"), FONT)
    doc.styles["Normal"].font.size = Pt(10)


def h(doc, text, level):
    p = doc.add_heading(text, level=level)
    return p


def para(doc, text, bold=False, size=10, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(9)
        # 헤더 배경 회색(무채색)
        shd = hdr[i]._tc.get_or_add_tcPr()
        el = shd.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): "DCDEE0"})
        shd.append(el)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
    return t


# ────────────────────────────────────────────────────────────────────────────
# 문서 1: 정책서
# ────────────────────────────────────────────────────────────────────────────
def build_policy():
    doc = Document()
    set_ko_font(doc)

    title = doc.add_heading("", level=0)
    r = title.add_run("도면검도시스템 정책서")
    r.font.size = Pt(22)
    sub = doc.add_paragraph()
    sr = sub.add_run("실외기 조합 매핑 구현 · 장비마스터 등록 정책")
    sr.italic = True
    sr.font.size = Pt(12)
    meta = doc.add_paragraph()
    meta.add_run("버전 v1.0  ·  작성일 2026-07-03  ·  대상: 개발팀/기획/QA  ·  근거: 프로젝트 코드·CLAUDE.md·doc/ 요구사항").font.size = Pt(9)
    doc.add_paragraph()

    # 1. 개요
    h(doc, "1. 문서 개요", 1)
    para(doc, "본 정책서는 LG전자 HVAC 도면검도시스템의 두 핵심 정책을 정의한다: (1) 유저 '생성' 단계의 실외기 조합 매핑, "
              "(2) 관리자 '장비마스터'의 등록·게시. 조합/호환 판단은 유저(생성) 단이 마스터의 게시(PUBLISHED) 데이터를 참조해 수행하며, "
              "장비마스터는 '등록만' 담당한다.")
    para(doc, "⚠️ 주의: doc/ 폴더 자료 중 일부는 최신화되지 않았을 수 있다. 문서 간/문서-코드 상충 또는 수치(단가·등급·스키마)가 "
              "애매할 경우 임의 추정하지 않고 담당자 확인 후 반영한다.", italic=True)

    # 2. 용어집
    h(doc, "2. 유비쿼터스 언어 (용어집)", 1)
    table(doc, ["용어", "코드/영문", "의미"], [
        ["실외기", "ODU / OutdoorUnit", "냉매를 실내기에 공급하는 실외 장비"],
        ["실내기", "IDU / IndoorUnit", "실(방)에 설치되는 냉난방 토출 장비"],
        ["조합비", "ComboRatio", "Σ(연결 실내기 냉방용량) ÷ 실외기 용량. 권장 0.5~1.3"],
        ["계열", "EnergySource", "EHP/GHP/AWHP/수냉식/Chiller. 호환의 기준"],
        ["조합/그룹", "OutdoorGroup", "실외기 1대 + 연결된 실내기들의 묶음(애그리거트)"],
        ["배정 플랜", "AssignmentPlan", "그룹들 + 미배정 풀의 조율자"],
        ["장비마스터", "EquipmentMaster", "장비 모델·스펙·단가 SSOT(관리자)"],
        ["게시 상태", "DRAFT/PUBLISHED/ARCHIVED", "외부 노출 게이트"],
        ["실", "Zone/Room", "도면에서 검출된 공간 단위(구 '방')"],
    ])

    # 3. 실외기 조합 매핑
    h(doc, "3. 실외기 조합 매핑 정책 (생성 컨텍스트)", 1)
    h(doc, "3.1 파이프라인", 2)
    para(doc, "도면 업로드 → [실 검출] → [실내기 자동 배치] → [실외기 선정] → [실외기 배치/조합 매핑] → [산출물 생성]. "
              "본 절은 '실외기 배치/조합 매핑' 단계의 정책을 정의한다.")
    h(doc, "3.2 도메인 모델", 2)
    table(doc, ["개념", "코드", "책임/불변식"], [
        ["실내기(엔티티)", "IndoorUnit", "id 동일성. 냉방용량(Capacity)·계열(EnergySource) 보유"],
        ["실외기(값객체)", "OutdoorUnit", "모델·계열·용량·최대연결수·단가(Price)·등급(EnergyGrade)"],
        ["조합(애그리거트 루트)", "OutdoorGroup", "그룹 내부 불변식 강제(계열/최대수/중복). 불변(immutable)"],
        ["배정 플랜(조율자)", "AssignmentPlan", "교차 불변식: 실내기는 정확히 한 곳에만 존재"],
    ])
    h(doc, "3.3 호환 규칙 (계열 일치)", 2)
    para(doc, "실외기와 실내기의 계열(EnergySource)이 일치해야 배정 가능하다. 교차 계열(예: GHP 실외기 ↔ EHP 실내기)은 "
              "SERIES_MISMATCH로 차단한다. 호환 판단 기준은 장비마스터의 제품군·계열 정보다.")
    h(doc, "3.4 조합비 규칙", 2)
    para(doc, "조합비 = Σ(연결 실내기 냉방용량) ÷ 실외기 용량. 권장 범위 0.5 ~ 1.3. 범위를 벗어나도 '배정 거부'가 아니라 "
              "'경고'(OVERLOADED/UNDERLOADED)로 표기한다 — 과부하 상태로도 배정은 가능해야 하기 때문.")
    h(doc, "3.5 최대 연결 수 (마스터 스펙 주입)", 2)
    para(doc, "실외기 모델별 최대 연결 실내기 수(maxConnections)는 장비마스터 카탈로그 스펙(OutdoorModelCatalog 포트)에서 주입된다. "
              "초과 배정은 MAX_CONNECTIONS로 차단한다. (하드코딩 금지 — 모델 스펙이 SSOT)")
    h(doc, "3.6 조작 유즈케이스", 2)
    table(doc, ["조작", "유즈케이스", "결과/규칙"], [
        ["실내기 배정/해제", "ReassignIndoorUnit", "그룹↔풀 이동. 계열 불일치는 거부(미저장), 성공 시 이벤트 발행"],
        ["실외기 모델 교체", "ReplaceOutdoorModel", "계열 변경 시 미호환 실내기를 풀로 방출"],
        ["그룹 분할", "SplitGroup", "실내기 절반을 같은 모델의 새 그룹으로"],
        ["그룹 추가", "AddGroup", "빈 실외기 그룹 추가"],
        ["그룹 삭제", "RemoveGroup", "연결 실내기는 풀로 반환"],
    ])
    h(doc, "3.7 예외·경고 정책", 2)
    table(doc, ["코드", "유형", "처리"], [
        ["SERIES_MISMATCH", "거부", "계열 불일치 배정 차단(모달 경고)"],
        ["MAX_CONNECTIONS", "거부", "최대 연결 수 초과 차단"],
        ["DUPLICATE", "거부", "이미 배정된 실내기 중복 차단"],
        ["OVERLOADED/UNDERLOADED", "경고", "조합비 범위 이탈. 배정은 허용, 표기만"],
        ["NotFound", "전파", "존재하지 않는 실내기/그룹 참조 오류"],
    ])
    h(doc, "3.8 UI 정책", 2)
    para(doc, "조합 리포트(상단 스트립)는 6개 KPI만 노출한다: 총 설치 용량 · 실외기 대수 · 실내기 배정 · 미배정 · 평균 조합비 · 과부하. "
              "(단가·효율등급은 리포트 KPI가 아님 — 산출물 장비일람표/모델 선정에서 다룸.) 조합 매핑은 팝업(모달)에서 드래그 배정/해제·"
              "교체·분할·삭제·조합비 실시간 계산·호환 경고로 수행한다. 색상은 무채색(회색조)만 사용한다.")

    # 4. 장비마스터 등록
    h(doc, "4. 장비마스터 등록 정책 (관리자)", 1)
    h(doc, "4.1 분류 체계 (4단 계층)", 2)
    table(doc, ["계층", "예시"], [
        ["대분류(역할)", "실외기 / 실내기 / 환기 / 판넬 / 시공자재 / 제어·통신"],
        ["중분류(제품군·계열)", "실외기: EHP·GHP·AWHP·수냉식·Chiller / 실내기: 1WAY·4WAY·덕트·스탠드·FCU"],
        ["시리즈", "Multi V Super 5, Multi V i, GHP Super III"],
        ["모델", "RPUW12BX9M, GPUW280C2S (마스터 레코드)"],
    ])
    h(doc, "4.2 스펙 정규화", 2)
    para(doc, "장비일람표 전 컬럼을 필수 정규화 필드로, 제품군 고유 항목은 JSONB 확장 스펙으로 저장한다. "
              "(근거: DB 설계 v2 스키마, 1,826 라벨 → 122 키 정규화 분석 보고서)")
    h(doc, "4.3 단가 정책", 2)
    para(doc, "단가는 product_prices(price_krw NUMERIC(14,0) 정수 원, price_with_vat_krw, price_type, effective_start/end_date)로 "
              "관리한다. 현행가 = effective_end_date IS NULL. 외부(생성/검도)는 게시뷰 v_published_product_prices의 현행가만 소비한다. "
              "도메인 값객체 Price가 정수·상한·유형 혼합 합산 금지 등 불변식을 강제한다.")
    h(doc, "4.4 게시 게이트", 2)
    para(doc, "DRAFT → PUBLISHED → ARCHIVED. 검도·생성은 PUBLISHED 데이터만 조회한다. 게시 전(DRAFT) 데이터는 외부 노출 금지.")
    h(doc, "4.5 등록 플로우 (단일 메뉴 내부 뷰)", 2)
    para(doc, "목록(필터·정렬·페이지네이션) / 상세·등록·수정(유효성 검증) / 엑셀 업로드·동기화 / 게시·단가. "
              "조합 매핑은 관리자가 아니라 유저(생성) 단에서 마스터 데이터를 참조해 수행한다.")
    h(doc, "4.6 컨텍스트 관계 (Customer/Supplier)", 2)
    para(doc, "장비마스터(상류) → 생성·검도(하류). 생성·검도는 마스터의 PUBLISHED 데이터만 소비하며 역방향 의존을 금지한다. "
              "생성 컨텍스트는 OutdoorModelCatalog 읽기 포트를 통해서만 마스터 스펙을 참조한다.")

    # 5. 아키텍처 원칙
    h(doc, "5. 아키텍처 원칙", 1)
    bullet(doc, "Clean Architecture: 의존성은 항상 안쪽(도메인)으로. domain은 상위 레이어 import 금지(eslint 게이트로 강제).")
    bullet(doc, "DDD: 값객체(불변·자기검증), 애그리거트 루트로만 상태 변경, 리포지토리 포트, 도메인 이벤트.")
    bullet(doc, "TDD: Red→Green→Refactor. 도메인/애플리케이션 커버리지 목표 90%+.")
    bullet(doc, "적대적 QA: 경계값·악의적 입력·실패 주입·동시성을 능동 탐색, 결함은 실패 테스트로 고정 후 수정.")
    bullet(doc, "품질 게이트: npm run validate (tsc --noEmit && eslint && vitest), Claude Hooks(가드/포맷/테스트).")

    # 6. 비고
    h(doc, "6. 비고 & 주의", 1)
    bullet(doc, "POC 플레이스홀더: ODU_CATALOG의 단가/등급/COP, maxConnections는 예시/보간값(실데이터 아님, 교체 예정).")
    bullet(doc, "무채색 규칙: UI는 유채색 금지(회색조만).")
    bullet(doc, "이름 규칙: 모든 사람 이름은 '홍길동'.")
    bullet(doc, "doc/ 신뢰도: 요구사항 정의서 v0.4 등 일부 최신 아닐 수 있음 — 애매 시 담당자 확인.")

    out = os.path.join(ROOT, "도면검도시스템_정책서_실외기조합매핑_장비마스터.docx")
    doc.save(out)
    return out


# ────────────────────────────────────────────────────────────────────────────
# WBS 데이터 (Epic → Task). 상태: 완료/진행/예정
# ────────────────────────────────────────────────────────────────────────────
# (id, 유형, 요약, 컴포넌트, 담당, 우선순위, SP, 상태, 완료조건)
WBS = [
    ("E1", "생성 - 실외기 조합 매핑", [
        ("E1-S1", "Task", "공유 값객체(Capacity/ComboRatio/EnergySource/ModelCode/Price/EnergyGrade)", "domain", "도메인", "High", 5, "완료", "불변·자기검증·경계 테스트 그린"),
        ("E1-S2", "Task", "OutdoorGroup 애그리거트(계열/최대수/중복 불변식)", "domain", "도메인", "High", 5, "완료", "assign/replace/split 불변식 테스트 그린"),
        ("E1-S3", "Task", "AssignmentPlan 조율자(교차 불변식: 정확히 한 곳)", "domain", "도메인", "High", 5, "완료", "reassign 유실·중복 없음 테스트"),
        ("E1-S4", "Task", "유즈케이스(reassign/replace/add/remove/split) + 포트/InMemory 리포지토리", "application", "BE", "High", 5, "완료", "규칙위반 거부·이벤트 발행 테스트"),
        ("E1-S5", "Task", "OutdoorModelCatalog 포트 + maxConnections 스펙 주입", "application", "BE", "Medium", 3, "완료", "모델 스펙 주입→불변식 구동 테스트"),
        ("E1-S6", "Task", "매핑 팝업 UI(드래그 배정/해제·교체·분할·삭제·조합비·경고)", "presentation", "FE", "High", 8, "완료", "동작 보존·회귀 없음"),
        ("E1-S7", "Task", "조합 리포트 6 KPI(설치용량/실외기대수/배정/미배정/평균조합비/과부하)", "presentation", "FE", "Medium", 3, "완료", "목업/요구사항 6 KPI 일치"),
        ("E1-S8", "Story", "실외기 자동 선정 로직(용량·조합비·효율 기준 최적 모델)", "domain", "BE", "High", 8, "예정", "부하 총량 기준 후보 조회·최적 선정·적정성 검증"),
        ("E1-S9", "Story", "조합 결과 도메인 영속화(서버 PlanRepository, SQS 재시도)", "infrastructure", "BE", "High", 8, "예정", "저장/복구·워커 재시도·동시성 처리"),
        ("E1-S10", "Story", "실외기 배치/구역 지정(CAD 좌표 산정·심볼 INSERT)", "domain", "BE", "Medium", 5, "예정", "지정 구역 내 배치 좌표 산정·블록 생성"),
    ]),
    ("E2", "생성 - 실 검출 & 도면 뷰어", [
        ("E2-S1", "Task", "SVG 뷰어(커서 줌·CTM 팬·단축키·플로팅 위젯·격자)", "presentation", "FE", "High", 8, "완료", "커서기준 줌·화면밖 팬 지속"),
        ("E2-S2", "Task", "실 다중선택(마퀴) + Figma 도구바(C/Z/H 모드)", "presentation", "FE", "Medium", 5, "완료", "영역선택·모드 전환"),
        ("E2-S3", "Task", "실내기 이동/회전/삭제 + 실 모서리 리사이즈(뷰어 로컬)", "presentation", "FE", "Medium", 5, "완료", "이동/회전(90°·15°)/삭제/리사이즈"),
        ("E2-S4", "Story", "실(방) 검출 AI 워커(DXF/PDF → 폴리곤·면적·라벨)", "worker", "BE", "High", 13, "예정", "도면→Room 목록, 검증"),
        ("E2-S5", "Story", "실내기 자동 배치(부하 계산 → 권장 모델·수량·타입)", "domain", "BE", "High", 8, "예정", "면적/부하 기준 수량·타입·좌표 산정"),
        ("E2-S6", "Story", "편집 결과 도메인 영속화(IndoorUnit 좌표/회전 유즈케이스)", "application", "BE", "Medium", 5, "예정", "배치 저장·복구"),
        ("E2-S7", "Story", "도면 업로드(PDF/DWG/DXF, S3 멀티파트, 진행률 WebSocket)", "infrastructure", "BE", "High", 8, "예정", "대용량 업로드·실시간 진행률"),
        ("E2-S8", "Task", "Viewer 상호작용 로직 커스텀 훅 분리(리팩터)", "presentation", "FE", "Low", 3, "예정", "Viewer<200줄, 훅 분리"),
    ]),
    ("E3", "장비마스터 (관리자)", [
        ("E3-S1", "Story", "4단 분류 체계 CRUD(대/중/시리즈/모델)", "admin", "FE/BE", "High", 8, "예정", "계층 등록·수정·조회"),
        ("E3-S2", "Story", "스펙 정규화(정규화 필드 + JSONB 확장)", "admin", "BE", "High", 8, "예정", "장비일람표 컬럼 정규화·확장 스펙 저장"),
        ("E3-S3", "Story", "엑셀 업로드/동기화(SPEC/CATALOG/PRICE)", "admin", "BE", "High", 8, "예정", "업로드·매핑·검증·동기화"),
        ("E3-S4", "Story", "게시 게이트(DRAFT→PUBLISHED→ARCHIVED)", "admin", "BE", "High", 5, "예정", "상태 전이·PUBLISHED만 외부 노출"),
        ("E3-S5", "Story", "단가 등록/수정(product_prices, 현행가/VAT)", "admin", "FE/BE", "High", 5, "예정", "단가 CRUD·현행가 게시뷰"),
        ("E3-S6", "Task", "목록 필터·정렬·페이지네이션", "admin", "FE", "Medium", 3, "예정", "유형/용량/모델 필터"),
    ]),
    ("E4", "검도 (Review)", [
        ("E4-S1", "Story", "도면-장비일람표 정합성 검증 로직", "domain", "BE", "High", 8, "예정", "정합/불일치 판정 규칙"),
        ("E4-S2", "Story", "검도 판정 결과 화면", "presentation", "FE", "Medium", 5, "예정", "판정 결과·불일치 목록"),
        ("E4-S3", "Story", "검도 리포트 PDF 생성(판정 결과 종합)", "worker", "BE", "Medium", 5, "예정", "PDF 자동 생성·이력"),
    ]),
    ("E5", "산출물 생성", [
        ("E5-S1", "Story", "장비일람표(실내기/실외기 탭) 생성", "domain", "BE", "High", 8, "예정", "탭별 일람표·필터·정렬"),
        ("E5-S2", "Task", "총 단가/집계(totalOutdoorPrice·gradeDistribution 활용)", "domain", "BE", "Medium", 3, "예정", "총액·유형별·미상 계상"),
        ("E5-S3", "Story", "배관 경로(Dijkstra) 길이 산출·집계", "domain", "BE", "Medium", 8, "예정", "실외기→실내기 최단경로·구역별 집계"),
        ("E5-S4", "Story", "설계 산출 종합 리포트(렌더·출력·저장)", "presentation", "FE/BE", "Medium", 5, "예정", "리포트 레이아웃·출력·저장"),
    ]),
    ("E6", "공통 / 아키텍처 / 인프라", [
        ("E6-S1", "Task", "Clean Arch 스캐폴딩 + 레이어 eslint 게이트", "infrastructure", "BE", "High", 3, "완료", "domain→상위 import 차단"),
        ("E6-S2", "Task", "TypeScript strict 전환", "infrastructure", "FE/BE", "High", 5, "완료", "tsc 0, 판별 유니온·포트 인터페이스"),
        ("E6-S3", "Task", "품질 게이트(validate: tsc/eslint/vitest) + Claude Hooks", "infrastructure", "BE", "Medium", 3, "완료", "Stop 훅 그린 게이트"),
        ("E6-S4", "Story", "인증/세션(로그인·토큰 갱신·권한·동시로그인)", "infrastructure", "BE", "High", 8, "예정", "세션 만료·403·토큰 갱신"),
        ("E6-S5", "Story", "CI 파이프라인(빌드·테스트·린트 게이트)", "infrastructure", "BE", "Medium", 3, "예정", "PR 게이트 자동화"),
    ]),
]


def build_wbs():
    doc = Document()
    set_ko_font(doc)

    title = doc.add_heading("", level=0)
    title.add_run("도면검도시스템 WBS · 지라 티켓 분류").font.size = Pt(20)
    doc.add_paragraph().add_run("개발 착수·지라 등록용 작업 분해 구조 (Epic → Story/Task)").italic = True
    doc.add_paragraph().add_run("버전 v1.0 · 2026-07-03 · 상태 기준: 프로젝트 현재 진행").font.size = Pt(9)
    doc.add_paragraph()

    h(doc, "사용법 (지라 등록)", 1)
    bullet(doc, "구조: Epic(대분류) → Story/Task(작업). Story는 사용자가치 단위, Task는 기술작업 단위.")
    bullet(doc, "지라 import: 동봉된 '지라_import_wbs.csv'를 Jira 'Import from CSV'로 등록. Epic Link 컬럼으로 상위 Epic에 연결.")
    bullet(doc, "필드: 컴포넌트(모듈), 담당(FE/BE/도메인/워커), 우선순위(High/Medium/Low), SP(스토리포인트), 상태, 완료조건(AC).")
    bullet(doc, "상태: 완료=구현·테스트 그린 / 진행=작업 중 / 예정=백로그.")

    h(doc, "우선순위·스토리포인트 기준", 1)
    table(doc, ["항목", "기준"], [
        ["High", "핵심 파이프라인·차단 요소·아키텍처 근간"],
        ["Medium", "주요 기능이나 대체·후속 가능"],
        ["Low", "리팩터·개선·편의"],
        ["SP(피보나치)", "1·2·3=소, 5=중, 8=대, 13=분할 권장(에픽화)"],
    ])

    # 요약 통계
    total = sum(len(tasks) for _, _, tasks in WBS)
    done = sum(1 for _, _, tasks in WBS for t in tasks if t[7] == "완료")
    h(doc, "요약", 1)
    para(doc, f"에픽 {len(WBS)}개 · 작업 {total}개 (완료 {done} / 예정·진행 {total - done}).")

    for eid, ename, tasks in WBS:
        h(doc, f"{eid}. {ename}", 1)
        rows = []
        for (tid, ttype, summ, comp, owner, prio, sp, status, ac) in tasks:
            rows.append([tid, ttype, summ, comp, owner, prio, sp, status, ac])
        table(doc, ["ID", "유형", "요약", "컴포넌트", "담당", "우선", "SP", "상태", "완료조건(AC)"], rows)

    out = os.path.join(ROOT, "도면검도시스템_WBS_지라티켓분류.docx")
    doc.save(out)
    return out


def build_csv():
    out = os.path.join(ROOT, "지라_import_wbs.csv")
    with open(out, "w", newline="", encoding="utf-8-sig") as f:
        w = csv.writer(f)
        w.writerow(["Issue Type", "Summary", "Epic Name", "Epic Link", "Component", "Assignee Group",
                    "Priority", "Story Points", "Status", "Labels", "Description"])
        for eid, ename, tasks in WBS:
            # Epic row
            w.writerow(["Epic", f"[{eid}] {ename}", ename, "", "", "", "High", "", "To Do", eid.lower(), f"에픽: {ename}"])
            for (tid, ttype, summ, comp, owner, prio, sp, status, ac) in tasks:
                jira_status = {"완료": "Done", "진행": "In Progress", "예정": "To Do"}.get(status, "To Do")
                w.writerow([ttype, f"[{tid}] {summ}", "", ename, comp, owner, prio, sp, jira_status,
                            f"{comp};{owner}", f"완료조건(AC): {ac}"])
    return out


if __name__ == "__main__":
    p1 = build_policy()
    p2 = build_wbs()
    p3 = build_csv()
    for p in (p1, p2, p3):
        print("생성:", os.path.basename(p))
